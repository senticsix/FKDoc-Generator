import platform
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from PyQt6.QtWidgets import QApplication

from ConfigManager import TEMPLATE_FILE
from WindowClass import FahrtkostenWindow


def iter_paragraphs(doc):
    """Yield all paragraphs: document body and all table cells."""
    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_in_paragraph(paragraph, old_text, new_text):
    """Replace text while preserving run formatting where possible.

    Word often splits text across multiple runs. First try a run-level
    replace (keeps all formatting). If the placeholder spans runs, merge
    the paragraph text into the first run (keeps that run's formatting).
    """
    if old_text not in paragraph.text:
        return

    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)

    if old_text not in paragraph.text:
        return

    merged = paragraph.text.replace(old_text, new_text)
    for run in paragraph.runs:
        run.text = ""

    if paragraph.runs:
        paragraph.runs[0].text = merged
    else:
        paragraph.add_run(merged)


def find_and_replace(doc, old_text, new_text):
    for paragraph in iter_paragraphs(doc):
        replace_in_paragraph(paragraph, old_text, new_text)


W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"


def _w14(tag):
    return f"{{{W14_NS}}}{tag}"


def _preceding_text(sdt):
    """Visible text before a content control within the same paragraph."""
    text = ""

    for sibling in sdt.getparent():
        if sibling is sdt:
            break
        text += "".join(t.text or "" for t in sibling.iter(qn("w:t")))

    return text


def _set_sdt_content_text(sdt, text):
    content = sdt.find(qn("w:sdtContent"))
    if content is None:
        return

    t_elements = list(content.iter(qn("w:t")))
    if t_elements:
        t_elements[0].text = text
        for t in t_elements[1:]:
            t.text = ""


def set_checkbox(doc, label, checked):
    """Toggle the Word checkbox content control right after the given label text."""
    for sdt in doc.element.body.iter(qn("w:sdt")):
        properties = sdt.find(qn("w:sdtPr"))
        if properties is None:
            continue

        checkbox = properties.find(_w14("checkbox"))
        if checkbox is None:
            continue

        if not _preceding_text(sdt).strip().endswith(label):
            continue

        checked_element = checkbox.find(_w14("checked"))
        if checked_element is None:
            checked_element = checkbox.makeelement(_w14("checked"), {})
            checkbox.append(checked_element)
        checked_element.set(_w14("val"), "1" if checked else "0")

        _set_sdt_content_text(sdt, "☒" if checked else "☐")
        return True

    print(f"Checkbox nach '{label}' nicht gefunden.")
    return False


def set_date_control(doc, anchor, date_text):
    """Fill the Word date-picker content control right after the anchor text."""
    try:
        date = datetime.strptime(date_text, "%d.%m.%Y")
    except ValueError:
        return False

    for sdt in doc.element.body.iter(qn("w:sdt")):
        properties = sdt.find(qn("w:sdtPr"))
        if properties is None:
            continue

        date_element = properties.find(qn("w:date"))
        if date_element is None:
            continue

        if not _preceding_text(sdt).strip().endswith(anchor):
            continue

        date_element.set(qn("w:fullDate"), date.strftime("%Y-%m-%dT00:00:00Z"))
        _set_sdt_content_text(sdt, date_text)
        return True

    print(f"Datumsfeld nach '{anchor}' nicht gefunden.")
    return False


def append_after_label(doc, label, text):
    """Append bold text to the paragraph that starts with the given label."""
    for paragraph in iter_paragraphs(doc):
        if paragraph.text.strip().startswith(label):
            run = paragraph.add_run(" " + text)
            run.bold = True
            return True

    return False


def build_signature_text(full_name):
    """Build a signature like "R. Remm" from a full name."""
    parts = full_name.split()

    if len(parts) >= 2:
        return f"{parts[0][0]}. {parts[-1]}"

    return full_name


def replace_text_with_signature(doc, placeholder, signature_text):
    """Replace the placeholder with the signature in a handwriting-style font."""
    for paragraph in iter_paragraphs(doc):
        if placeholder in paragraph.text:
            replace_in_paragraph(paragraph, placeholder, "")

            run = paragraph.add_run(signature_text)
            run.font.name = "Brush Script MT"
            run.font.size = Pt(22)


def convert_docx_to_pdf(docx_path):
    """Convert DOCX to PDF using MS Word, depending on the platform."""
    system = platform.system()

    if system == "Windows":
        return _convert_docx_to_pdf_windows(docx_path)

    if system == "Darwin":
        return _convert_docx_to_pdf_mac(docx_path)

    print(f"PDF conversion is not supported on {system}.")
    return None


def _convert_docx_to_pdf_windows(docx_path):
    """Convert DOCX to PDF using MS Word COM on Windows."""
    word = None

    try:
        import win32com.client

        docx_path = Path(docx_path).absolute()
        pdf_path = docx_path.with_suffix(".pdf")

        print(f"Converting: {docx_path} -> {pdf_path}")

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        doc = word.Documents.Open(str(docx_path))

        # FileFormat=17 is wdFormatPDF
        doc.SaveAs(str(pdf_path), FileFormat=17)
        doc.Close()

        print(f"PDF created successfully: {pdf_path}")
        return str(pdf_path)
    except ImportError as e:
        print(f"pywin32 not installed: {e}")
        return None
    except Exception as e:
        print(f"Windows PDF conversion failed: {e}")
        import traceback

        traceback.print_exc()
        return None
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def _applescript_escape(path):
    return str(path).replace("\\", "\\\\").replace('"', '\\"')


def _convert_docx_to_pdf_mac(docx_path):
    """Convert DOCX to PDF using MS Word via AppleScript on macOS."""
    try:
        docx_path = Path(docx_path).absolute()
        pdf_path = docx_path.with_suffix(".pdf")

        print(f"Converting: {docx_path} -> {pdf_path}")

        # Only closes the converted document - does not quit Word, in case
        # the user has other documents open.
        applescript = f'''
        tell application "Microsoft Word"
            open POSIX file "{_applescript_escape(docx_path)}"
            set theDoc to active document
            save as theDoc file name "{_applescript_escape(pdf_path)}" file format format PDF
            close theDoc saving no
        end tell
        '''

        result = subprocess.run(
            ["osascript", "-e", applescript],
            capture_output=True,
            text=True,
            timeout=60,
        )

        if result.returncode == 0 and pdf_path.exists():
            print(f"PDF created successfully: {pdf_path}")
            return str(pdf_path)

        print("PDF file not created.")
        print(f"stderr: {result.stderr}")
        print(f"stdout: {result.stdout}")
        return None

    except subprocess.TimeoutExpired:
        print("PDF conversion timed out")
        return None
    except Exception as e:
        print(f"macOS PDF conversion failed: {e}")
        import traceback

        traceback.print_exc()
        return None


def organize_pdf(pdf_path, date_from, date_to, config):
    """Move the PDF into the configured root folder, ordered by date.

    Structure: <root>/<YYYY>/<MM>/Fahrtkostenerstattung_<von>_bis_<bis>.pdf
    Without a configured root folder the PDF stays next to the DOCX.
    """
    root = str(config.get("pdf-root", "")).strip()

    if not root:
        return pdf_path

    try:
        d_from = datetime.strptime(date_from, "%d.%m.%Y")
        d_to = datetime.strptime(date_to, "%d.%m.%Y")

        target_dir = Path(root) / f"{d_from:%Y}" / f"{d_from:%m}"
        target_dir.mkdir(parents=True, exist_ok=True)

        target = target_dir / (
            f"Fahrtkostenerstattung_{d_from:%Y-%m-%d}_bis_{d_to:%Y-%m-%d}.pdf"
        )

        shutil.move(str(pdf_path), str(target))
        print(f"PDF abgelegt unter: {target}")
        return str(target)
    except Exception as e:
        print(f"PDF-Ablage im Root-Ordner fehlgeschlagen: {e}")
        return pdf_path


def generate_document(
    hin,
    back,
    hin1=None,
    back1=None,
    km="",
    config=None,
    krank=None,
):
    if config is None:
        config = {}

    curr_date = datetime.today().strftime("%d.%m.%Y")

    if hin1 is None:
        datum = hin + " - " + back
    else:
        datum = hin + " - " + back1

    default_output = Path.home() / "Desktop" / "Fahrtkostenerstattung_Familienheimfahrt.docx"
    output_path = config.get("output_path") or str(default_output)

    document = Document(str(TEMPLATE_FILE))

    find_and_replace(document, "{DATUM}", datum)
    find_and_replace(document, "{DATUM_HIN}", hin)
    find_and_replace(document, "{DATUM_BACK}", back)
    find_and_replace(document, "{KM}", km)
    find_and_replace(document, "{SIG_DATE}", curr_date)

    find_and_replace(document, "{NAME}", config.get("name", ""))
    find_and_replace(document, "{DATUM_ERSTANR}", config.get("date-init", ""))
    # Ausbildungsdauer: als Text-Platzhalter UND als Word-Datumssteuerelemente
    find_and_replace(document, "{DATUM_AUSB_ANF}", config.get("date-ausb-anf", ""))
    find_and_replace(document, "{DATUM_AUSB_ENDE}", config.get("date-ausb-ende", ""))
    set_date_control(document, "vom", config.get("date-ausb-anf", ""))
    set_date_control(document, "bis", config.get("date-ausb-ende", ""))

    # Krankmeldung: Ja/Nein-Kontrollkästchen und ggf. Zeitraum
    has_sick_note = krank is not None
    set_checkbox(document, "Ja:", has_sick_note)
    set_checkbox(document, "Nein:", not has_sick_note)

    if has_sick_note:
        append_after_label(
            document,
            "Zeitraum der Krankmeldung",
            f"{krank[0]} – {krank[1]}",
        )
    signature = build_signature_text(config.get("name", ""))
    replace_text_with_signature(document, "{SIGNATURE}", signature)

    if hin1 is None:
        find_and_replace(document, "{DATUM_HIN1}", "")
        find_and_replace(document, "{DATUM_BACK1}", "")
    else:
        find_and_replace(document, "{DATUM_HIN1}", hin1)
        find_and_replace(document, "{DATUM_BACK1}", back1)

    document.save(output_path)
    print(f"DOCX saved to: {output_path}")

    pdf_path = convert_docx_to_pdf(output_path)

    if pdf_path:
        pdf_path = organize_pdf(pdf_path, hin, back1 or back, config)

    return pdf_path


def main():
    app = QApplication(sys.argv)

    window = FahrtkostenWindow(generate_document)
    window.show()

    sys.exit(app.exec())


if __name__ == "__main__":
    main()
